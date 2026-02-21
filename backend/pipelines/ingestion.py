"""
backend/pipelines/ingestion.py
Handles PDF/DOCX ingestion: extraction → chunking → embedding → FAISS storage.
"""

import os
import shutil
from pathlib import Path
from typing import List, Tuple

import fitz  # PyMuPDF
import pdfplumber
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain.schema import Document

from backend.utils.config import config


# ─────────────────────────────────────────────
# Embedding model (singleton)
# ─────────────────────────────────────────────
_embeddings = None


def get_embeddings() -> HuggingFaceEmbeddings:
    global _embeddings
    if _embeddings is None:
        print(f"[Ingestion] Loading embedding model: {config.EMBEDDING_MODEL}")
        _embeddings = HuggingFaceEmbeddings(
            model_name=config.EMBEDDING_MODEL,
            model_kwargs={"device": "cpu"},
            encode_kwargs={"normalize_embeddings": True},
        )
    return _embeddings


# ─────────────────────────────────────────────
# Extraction
# ─────────────────────────────────────────────
def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF with pdfplumber fallback."""
    text = ""
    try:
        doc = fitz.open(file_path)
        for page in doc:
            text += page.get_text("text") + "\n"
        doc.close()
    except Exception as e:
        print(f"[Ingestion] PyMuPDF failed ({e}), trying pdfplumber...")
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e2:
            raise RuntimeError(f"Both PDF extractors failed: {e2}")
    return text.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX."""
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table content
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def extract_text(file_path: str) -> Tuple[str, str]:
    """
    Auto-detect file type and extract text.
    Returns (text, file_type).
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path), "pdf"
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path), "docx"
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX supported.")


# ─────────────────────────────────────────────
# Chunking
# ─────────────────────────────────────────────
def chunk_text(text: str, file_name: str) -> List[Document]:
    """Split text into overlapping chunks with metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.CHUNK_SIZE,
        chunk_overlap=config.CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )
    chunks = splitter.create_documents(
        [text],
        metadatas=[{"source": file_name, "chunk_id": 0}],
    )
    # Fix chunk_id after splitting
    for i, chunk in enumerate(chunks):
        chunk.metadata["chunk_id"] = i
        chunk.metadata["total_chunks"] = len(chunks)
    print(f"[Ingestion] Created {len(chunks)} chunks from '{file_name}'")
    return chunks


# ─────────────────────────────────────────────
# Vector Store
# ─────────────────────────────────────────────
def build_vectorstore(chunks: List[Document], store_path: str) -> FAISS:
    """Embed chunks and build/save FAISS index."""
    os.makedirs(store_path, exist_ok=True)
    embeddings = get_embeddings()
    print(f"[Ingestion] Building FAISS index with {len(chunks)} chunks...")
    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local(store_path)
    print(f"[Ingestion] FAISS index saved to: {store_path}")
    return vectorstore


def load_vectorstore(store_path: str) -> FAISS:
    """Load an existing FAISS index from disk."""
    embeddings = get_embeddings()
    vectorstore = FAISS.load_local(
        store_path, embeddings, allow_dangerous_deserialization=True
    )
    print(f"[Ingestion] FAISS index loaded from: {store_path}")
    return vectorstore


def append_to_vectorstore(
    new_chunks: List[Document], store_path: str
) -> FAISS:
    """Append new documents to an existing FAISS index."""
    if os.path.exists(os.path.join(store_path, "index.faiss")):
        vectorstore = load_vectorstore(store_path)
        vectorstore.add_documents(new_chunks)
    else:
        vectorstore = build_vectorstore(new_chunks, store_path)
    vectorstore.save_local(store_path)
    return vectorstore


# ─────────────────────────────────────────────
# Main Ingestion Entry Point
# ─────────────────────────────────────────────
def ingest_document(file_path: str, reset_store: bool = False) -> dict:
    """
    Full ingestion pipeline:
      file → extract → chunk → embed → FAISS

    Args:
        file_path: Path to the uploaded PDF/DOCX.
        reset_store: If True, wipe existing vectorstore before ingestion.

    Returns:
        dict with metadata about the ingestion.
    """
    file_name = Path(file_path).name
    store_path = config.VECTORSTORE_PATH

    # Optional reset
    if reset_store and os.path.exists(store_path):
        shutil.rmtree(store_path)
        print(f"[Ingestion] Cleared existing vectorstore at {store_path}")

    # Step 1: Extract
    print(f"[Ingestion] Extracting text from: {file_name}")
    text, file_type = extract_text(file_path)

    if not text or len(text) < 50:
        raise ValueError("Could not extract meaningful text from the document.")

    # Step 2: Chunk
    chunks = chunk_text(text, file_name)

    # Step 3: Embed + Store
    vectorstore = append_to_vectorstore(chunks, store_path)

    return {
        "status": "success",
        "file_name": file_name,
        "file_type": file_type,
        "num_chunks": len(chunks),
        "char_count": len(text),
        "vectorstore_path": store_path,
        "raw_text": text,  # kept for summarization pipeline
    }
